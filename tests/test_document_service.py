import hashlib
import asyncio
from pathlib import Path
import tempfile

import pytest
from docx import Document as DocxDocument
from reportlab.pdfgen import canvas

from backend.services.document_service import DocumentService
from backend.models.document import UploadedDocument, DataModule, PublicationModule
from backend.models.base import DMTypeEnum, SecurityLevel
import zipfile
import os
import types
from backend.ai_providers.provider_factory import ProviderFactory


def create_docx(path: Path, text: str):
    doc = DocxDocument()
    doc.add_paragraph(text)
    doc.save(path)


def create_pdf(path: Path):
    c = canvas.Canvas(str(path))
    c.drawString(100, 750, "Test PDF")
    c.showPage()
    c.save()


def test_extract_docx_text():
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = Path(tmpdir) / "sample.docx"
        text = "Hello World"
        create_docx(docx_path, text)

        service = DocumentService(upload_path=tmpdir)
        extracted = asyncio.run(service._extract_docx_text(docx_path))
        assert text in extracted


def test_extract_pdf_images():
    with tempfile.TemporaryDirectory() as tmpdir:
        pdf_path = Path(tmpdir) / "sample.pdf"
        create_pdf(pdf_path)

        with open(pdf_path, "rb") as f:
            data = f.read()
        sha = hashlib.sha256(data).hexdigest()

        doc = UploadedDocument(
            filename="sample.pdf",
            file_path=str(pdf_path),
            mime_type="application/pdf",
            file_size=len(data),
            sha256_hash=sha,
            metadata={},
        )

        service = DocumentService(upload_path=tmpdir)
        images = asyncio.run(service._extract_pdf_images(doc))
        assert isinstance(images, list)
        if images:
            for icn in images:
                assert Path(icn.file_path).exists()
                assert icn.width > 0 and icn.height > 0


class FakeCursor:
    def __init__(self, docs):
        self.docs = docs

    async def to_list(self, _):
        return self.docs


class FakeCollection:
    def __init__(self, docs):
        self.docs = docs

    def find(self, query):
        dms = [dm.dict() for dm in self.docs if dm.dmc in query["dmc"]["$in"]]
        return FakeCursor(dms)


class FakeDB:
    def __init__(self, docs):
        self.data_modules = FakeCollection(docs)


def test_publish_publication_module(tmp_path):
    dm1 = DataModule(
        dmc="DMC-TEST-0001",
        title="Test",
        dm_type=DMTypeEnum.GEN,
        info_variant="00",
        content="Hello",
        source_document_id="doc1",
    )
    dm2 = DataModule(
        dmc="DMC-TEST-0002",
        title="Test",
        dm_type=DMTypeEnum.GEN,
        info_variant="01",
        content="World",
        source_document_id="doc1",
    )

    pm = PublicationModule(pm_code="PM1", title="PM", dm_list=[dm1.dmc, dm2.dmc])
    service = DocumentService(upload_path=tmp_path)
    db = FakeDB([dm1, dm2])

    result = asyncio.run(
        service.publish_publication_module(
            pm, db, formats=["xml", "html", "pdf"], variants=["00", "01"]
        )
    )

    package = result["package"]
    assert result["errors"] == []
    assert package.exists()
    with zipfile.ZipFile(package) as z:
        names = z.namelist()
        assert f"{dm1.dmc}_{dm1.info_variant}.xml" in names
        assert f"{dm1.dmc}_{dm1.info_variant}.html" in names
        assert f"{dm1.dmc}_{dm1.info_variant}.pdf" in names


def test_process_document_carries_security_and_warnings(tmp_path):
    text = "WARNING: Hot surface\nCAUTION: Wear gloves\nStep 1"
    file_path = tmp_path / "s.txt"
    file_path.write_text(text)
    sha = hashlib.sha256(text.encode()).hexdigest()
    doc = UploadedDocument(
        filename="s.txt",
        file_path=str(file_path),
        mime_type="text/plain",
        file_size=len(text),
        sha256_hash=sha,
        security_level=SecurityLevel.SECRET,
        metadata={},
    )

    class DummyProvider:
        async def classify_document(self, request):
            return types.SimpleNamespace(result={"dm_type": "GEN", "title": "T"})

        async def extract_structured_data(self, request):
            return types.SimpleNamespace(result={"references": [], "warnings": ["Hot surface"], "cautions": ["Wear gloves"]})

        async def rewrite_to_ste(self, request):
            return types.SimpleNamespace(result={"rewritten_text": request.text, "ste_score": 1.0})

    service = DocumentService(upload_path=tmp_path)
    orig_factory = ProviderFactory.create_text_provider
    ProviderFactory.create_text_provider = lambda: DummyProvider()
    try:
        modules = asyncio.run(service.process_document_with_ai(doc, text))
    finally:
        ProviderFactory.create_text_provider = orig_factory
    assert modules
    for m in modules:
        assert m.security_level == SecurityLevel.SECRET
        assert "<warning>" in m.content
        assert "<caution>" in m.content
