"""Main FastAPI server for Aquila S1000D-AI system."""

import base64
import io
import json
import logging
import os
import tempfile
import uuid
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

import xmlschema
import yaml
from dotenv import load_dotenv
from fastapi import APIRouter, Depends, FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse
from lxml import etree
from motor.motor_asyncio import AsyncIOMotorClient

from .ai_providers.provider_factory import ProviderFactory
from .models.base import ProviderEnum, SettingsModel, ValidationStatus
from .models.document import (
    ICN,
    DataModule,
    ProcessingTask,
    PublicationModule,
    UploadedDocument,
)
from .services.document_service import DocumentService

# Load environment variables
ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / ".env")

# MongoDB connection
mongo_url = os.environ["MONGO_URL"]
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ["DB_NAME"]]

# Initialize document service
document_service = DocumentService()
BREX_RULES_PATH = ROOT_DIR / "schemas" / "brex_rules.yaml"
BREX_RULES = {}
if BREX_RULES_PATH.exists():
    with open(BREX_RULES_PATH, "r", encoding="utf-8") as f:
        BREX_RULES = yaml.safe_load(f) or {}

# Create FastAPI app
app = FastAPI(
    title="Aquila S1000D-AI API",
    description="AI-powered technical documentation processing system",
    version="1.0.0",
)

# Create API router
api_router = APIRouter(prefix="/api")

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure logging
logging.basicConfig(
    level=logging.INFO, format="%(asctime)s - %(name)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger(__name__)

# Global settings
system_settings = SettingsModel()

# API Endpoints


@api_router.get("/")
async def root():
    """Root endpoint."""
    return {"message": "Aquila S1000D-AI API", "version": "1.0.0"}


@api_router.get("/health")
async def health_check():
    """Health check endpoint."""
    provider_config = ProviderFactory.validate_provider_config()
    return {
        "status": "healthy",
        "timestamp": datetime.utcnow().isoformat(),
        "providers": provider_config,
    }


# Settings endpoints
@api_router.get("/settings")
async def get_settings():
    """Get current system settings."""
    return system_settings.dict()


@api_router.post("/settings")
async def update_settings(settings: SettingsModel):
    """Update system settings."""
    global system_settings
    system_settings = settings

    # Update environment variables
    os.environ["TEXT_PROVIDER"] = settings.text_provider.value
    os.environ["VISION_PROVIDER"] = settings.vision_provider.value
    os.environ["TEXT_MODEL"] = settings.text_model
    os.environ["VISION_MODEL"] = settings.vision_model

    return {
        "message": "Settings updated successfully",
        "settings": system_settings.dict(),
    }


# AI Provider endpoints
@api_router.get("/providers")
async def get_providers():
    """Get available AI providers and their status."""
    return {
        "available": ProviderFactory.get_available_providers(),
        "current": {
            "text": os.environ.get("TEXT_PROVIDER", "openai"),
            "vision": os.environ.get("VISION_PROVIDER", "openai"),
            "text_model": os.environ.get("TEXT_MODEL", "gpt-4o-mini"),
            "vision_model": os.environ.get("VISION_MODEL", "gpt-4o-mini"),
        },
        "config": ProviderFactory.validate_provider_config(),
    }


@api_router.post("/providers/set")
async def set_providers(
    text_provider: str,
    vision_provider: str,
    text_model: Optional[str] = None,
    vision_model: Optional[str] = None,
):
    """Set AI providers."""
    try:
        # Validate providers
        available = ProviderFactory.get_available_providers()
        if text_provider not in available["text"]:
            raise HTTPException(400, f"Invalid text provider: {text_provider}")
        if vision_provider not in available["vision"]:
            raise HTTPException(400, f"Invalid vision provider: {vision_provider}")

        # Update environment
        os.environ["TEXT_PROVIDER"] = text_provider
        os.environ["VISION_PROVIDER"] = vision_provider
        if text_model:
            os.environ["TEXT_MODEL"] = text_model
        if vision_model:
            os.environ["VISION_MODEL"] = vision_model

        # Update system settings
        system_settings.text_provider = ProviderEnum(text_provider)
        system_settings.vision_provider = ProviderEnum(vision_provider)
        if text_model:
            system_settings.text_model = text_model
        if vision_model:
            system_settings.vision_model = vision_model

        return {
            "message": "Providers updated successfully",
            "text_provider": text_provider,
            "vision_provider": vision_provider,
            "text_model": text_model or os.environ.get("TEXT_MODEL", "gpt-4o-mini"),
            "vision_model": vision_model
            or os.environ.get("VISION_MODEL", "gpt-4o-mini"),
        }
    except Exception as e:
        raise HTTPException(500, f"Error updating providers: {str(e)}")


# Document upload endpoints
@api_router.post("/documents/upload")
async def upload_document(file: UploadFile = File(...)):
    """Upload a document for processing."""
    try:
        # Read file data
        file_data = await file.read()

        # Upload document
        document = await document_service.upload_document(
            file_data=file_data, filename=file.filename, mime_type=file.content_type
        )

        # Store in database
        await db.documents.insert_one(document.dict())

        return {
            "message": "Document uploaded successfully",
            "document_id": document.id,
            "filename": document.filename,
            "size": document.file_size,
        }
    except Exception as e:
        logger.error(f"Error uploading document: {str(e)}")
        raise HTTPException(500, f"Error uploading document: {str(e)}")


@api_router.get("/documents")
async def get_documents():
    """Get all uploaded documents."""
    try:
        documents = await db.documents.find().to_list(1000)
        return [UploadedDocument(**doc) for doc in documents]
    except Exception as e:
        logger.error(f"Error fetching documents: {str(e)}")
        raise HTTPException(500, f"Error fetching documents: {str(e)}")


@api_router.get("/documents/{document_id}")
async def get_document(document_id: str):
    """Get a specific document."""
    try:
        document = await db.documents.find_one({"id": document_id})
        if not document:
            raise HTTPException(404, "Document not found")
        return UploadedDocument(**document)
    except Exception as e:
        logger.error(f"Error fetching document: {str(e)}")
        raise HTTPException(500, f"Error fetching document: {str(e)}")


@api_router.post("/documents/{document_id}/process")
async def process_document(document_id: str):
    """Process a document to create data modules."""
    try:
        # Get document
        doc_data = await db.documents.find_one({"id": document_id})
        if not doc_data:
            raise HTTPException(404, "Document not found")

        document = UploadedDocument(**doc_data)

        # Extract images first so we can embed references in the text
        images = await document_service.extract_images_from_document(document)

        # Extract text content with image references
        text_content = await document_service.extract_text_from_document(
            document, images
        )

        # Process images with AI
        processed_images = []
        for image in images:
            processed_image = await document_service.process_image_with_ai(image)
            processed_images.append(processed_image)
            # Store ICN in database
            await db.icns.insert_one(processed_image.dict())

        # Process document with AI
        data_modules = await document_service.process_document_with_ai(
            document, text_content, processed_images
        )

        # Store data modules in database
        stored_modules = []
        for dm in data_modules:
            await db.data_modules.insert_one(dm.dict())
            stored_modules.append(dm)

        # Update document status
        await db.documents.update_one(
            {"id": document_id},
            {
                "$set": {
                    "processing_status": "completed",
                    "updated_at": datetime.utcnow(),
                }
            },
        )

        return {
            "message": "Document processed successfully",
            "document_id": document_id,
            "data_modules": len(stored_modules),
            "images": len(processed_images),
            "modules": [dm.dict() for dm in stored_modules],
        }
    except Exception as e:
        logger.error(f"Error processing document: {str(e)}")
        raise HTTPException(500, f"Error processing document: {str(e)}")


# Data Module endpoints
@api_router.get("/data-modules")
async def get_data_modules():
    """Get all data modules."""
    try:
        modules = await db.data_modules.find().to_list(1000)
        return [DataModule(**module) for module in modules]
    except Exception as e:
        logger.error(f"Error fetching data modules: {str(e)}")
        raise HTTPException(500, f"Error fetching data modules: {str(e)}")


@api_router.get("/data-modules/{dmc}")
async def get_data_module(dmc: str):
    """Get a specific data module."""
    try:
        module = await db.data_modules.find_one({"dmc": dmc})
        if not module:
            raise HTTPException(404, "Data module not found")
        return DataModule(**module)
    except Exception as e:
        logger.error(f"Error fetching data module: {str(e)}")
        raise HTTPException(500, f"Error fetching data module: {str(e)}")


@api_router.put("/data-modules/{dmc}")
async def update_data_module(dmc: str, module_data: Dict[str, Any]):
    """Update a data module."""
    try:
        # Update module in database
        result = await db.data_modules.update_one(
            {"dmc": dmc}, {"$set": {**module_data, "updated_at": datetime.utcnow()}}
        )

        if result.matched_count == 0:
            raise HTTPException(404, "Data module not found")

        return {"message": "Data module updated successfully"}
    except Exception as e:
        logger.error(f"Error updating data module: {str(e)}")
        raise HTTPException(500, f"Error updating data module: {str(e)}")


@api_router.delete("/data-modules/{dmc}")
async def delete_data_module(dmc: str):
    """Delete a data module."""
    try:
        result = await db.data_modules.delete_one({"dmc": dmc})
        if result.deleted_count == 0:
            raise HTTPException(404, "Data module not found")
        return {"message": "Data module deleted successfully"}
    except Exception as e:
        logger.error(f"Error deleting data module: {str(e)}")
        raise HTTPException(500, f"Error deleting data module: {str(e)}")


# ICN endpoints
@api_router.get("/icns")
async def get_icns():
    """Get all ICNs."""
    try:
        icns = await db.icns.find().to_list(1000)
        return [ICN(**icn) for icn in icns]
    except Exception as e:
        logger.error(f"Error fetching ICNs: {str(e)}")
        raise HTTPException(500, f"Error fetching ICNs: {str(e)}")


@api_router.get("/icns/{icn_id}")
async def get_icn(icn_id: str):
    """Get a specific ICN."""
    try:
        icn = await db.icns.find_one({"icn_id": icn_id})
        if not icn:
            raise HTTPException(404, "ICN not found")
        return ICN(**icn)
    except Exception as e:
        logger.error(f"Error fetching ICN: {str(e)}")
        raise HTTPException(500, f"Error fetching ICN: {str(e)}")


@api_router.get("/icns/{icn_id}/image")
async def get_icn_image(icn_id: str):
    """Get ICN image file."""
    try:
        icn = await db.icns.find_one({"icn_id": icn_id})
        if not icn:
            raise HTTPException(404, "ICN not found")

        icn_obj = ICN(**icn)
        return FileResponse(
            icn_obj.file_path, media_type=icn_obj.mime_type, filename=icn_obj.filename
        )
    except Exception as e:
        logger.error(f"Error fetching ICN image: {str(e)}")
        raise HTTPException(500, f"Error fetching ICN image: {str(e)}")


@api_router.put("/icns/{icn_id}")
async def update_icn(icn_id: str, icn_data: Dict[str, Any]):
    """Update an ICN."""
    try:
        result = await db.icns.update_one(
            {"icn_id": icn_id}, {"$set": {**icn_data, "updated_at": datetime.utcnow()}}
        )

        if result.matched_count == 0:
            raise HTTPException(404, "ICN not found")

        return {"message": "ICN updated successfully"}
    except Exception as e:
        logger.error(f"Error updating ICN: {str(e)}")
        raise HTTPException(500, f"Error updating ICN: {str(e)}")


# Validation endpoints
@api_router.post("/validate/{dmc}")
async def validate_data_module(dmc: str):
    """Validate a data module."""
    try:
        module = await db.data_modules.find_one({"dmc": dmc})
        if not module:
            raise HTTPException(404, "Data module not found")

        validation_errors = []
        validation_status = ValidationStatus.GREEN

        if not module.get("title"):
            validation_errors.append("Title is required")
            validation_status = ValidationStatus.RED

        if not module.get("content"):
            validation_errors.append("Content is required")
            validation_status = ValidationStatus.RED

        # XSD validation
        xml_valid = False
        schema_path = Path(__file__).parent / "schemas" / "data_module.xsd"
        if module.get("xml_content"):
            try:
                schema = xmlschema.XMLSchema(str(schema_path))
                schema.validate(module["xml_content"])
                xml_valid = True
            except Exception as exc:
                validation_errors.append(f"XSD: {str(exc)}")
                validation_status = ValidationStatus.RED

        # BREX rules
        if (
            BREX_RULES.get("min_ste_score")
            and module.get("ste_score", 0) < BREX_RULES["min_ste_score"]
        ):
            validation_errors.append("STE score below minimum")
            if validation_status == ValidationStatus.GREEN:
                validation_status = ValidationStatus.AMBER

        if (
            BREX_RULES.get("max_title_length")
            and len(module.get("title", "")) > BREX_RULES["max_title_length"]
        ):
            validation_errors.append("Title too long")
            validation_status = ValidationStatus.RED

        # Update module validation status
        await db.data_modules.update_one(
            {"dmc": dmc},
            {
                "$set": {
                    "validation_status": validation_status.value,
                    "validation_errors": validation_errors,
                    "xsd_valid": xml_valid,
                    "updated_at": datetime.utcnow(),
                }
            },
        )

        return {
            "dmc": dmc,
            "status": validation_status.value,
            "errors": validation_errors,
        }
    except Exception as e:
        logger.error(f"Error validating data module: {str(e)}")
        raise HTTPException(500, f"Error validating data module: {str(e)}")


# Publication Module endpoints
@api_router.get("/publication-modules")
async def get_publication_modules():
    """Get all publication modules."""
    try:
        pms = await db.publication_modules.find().to_list(1000)
        return [PublicationModule(**pm) for pm in pms]
    except Exception as e:
        logger.error(f"Error fetching publication modules: {str(e)}")
        raise HTTPException(500, f"Error fetching publication modules: {str(e)}")


@api_router.post("/publication-modules")
async def create_publication_module(pm_data: Dict[str, Any]):
    """Create a new publication module."""
    try:
        pm = PublicationModule(**pm_data)
        await db.publication_modules.insert_one(pm.dict())
        return {
            "message": "Publication module created successfully",
            "pm_code": pm.pm_code,
        }
    except Exception as e:
        logger.error(f"Error creating publication module: {str(e)}")
        raise HTTPException(500, f"Error creating publication module: {str(e)}")


@api_router.get("/publication-modules/{pm_code}")
async def get_publication_module(pm_code: str):
    """Get a specific publication module."""
    try:
        pm = await db.publication_modules.find_one({"pm_code": pm_code})
        if not pm:
            raise HTTPException(404, "Publication module not found")
        return PublicationModule(**pm)
    except Exception as e:
        logger.error(f"Error fetching publication module: {str(e)}")
        raise HTTPException(500, f"Error fetching publication module: {str(e)}")


@api_router.post("/publication-modules/{pm_code}/publish")
async def publish_publication_module(pm_code: str, publish_options: Dict[str, Any]):
    """Publish a publication module."""
    try:
        pm = await db.publication_modules.find_one({"pm_code": pm_code})
        if not pm:
            raise HTTPException(404, "Publication module not found")

        temp_dir = tempfile.mkdtemp(prefix="publish_")
        package_path = Path(temp_dir) / f"{pm_code}.zip"

        with zipfile.ZipFile(package_path, "w") as pkg:
            for dmc in pm.get("dm_list", []):
                dm = await db.data_modules.find_one({"dmc": dmc})
                if not dm:
                    continue

                xml_content = dm.get("xml_content")
                if not xml_content:
                    root = etree.Element("dm")
                    title_el = etree.SubElement(root, "title")
                    title_el.text = dm.get("title", "")
                    body_el = etree.SubElement(root, "body")
                    body_el.text = dm.get("content", "")
                    xml_content = etree.tostring(
                        root, pretty_print=True, encoding="utf-8"
                    )

                base_path = Path(temp_dir) / dmc
                base_path.parent.mkdir(parents=True, exist_ok=True)

                if "xml" in publish_options.get("formats", []):
                    xml_file = base_path.with_suffix(".xml")
                    with open(xml_file, "wb") as f:
                        f.write(xml_content)
                    pkg.write(xml_file, arcname=xml_file.name)

                if "html" in publish_options.get("formats", []):
                    html_file = base_path.with_suffix(".html")
                    body = dm.get("content", "").replace("\n", "<br/>")
                    html = f"<h1>{dm.get('title','')}</h1><p>{body}</p>"
                    with open(html_file, "w", encoding="utf-8") as f:
                        f.write(html)
                    pkg.write(html_file, arcname=html_file.name)

                if "pdf" in publish_options.get("formats", []):
                    from reportlab.lib.pagesizes import letter
                    from reportlab.pdfgen import canvas

                    pdf_file = base_path.with_suffix(".pdf")
                    c = canvas.Canvas(str(pdf_file), pagesize=letter)
                    c.setFont("Helvetica", 12)
                    c.drawString(72, 720, dm.get("title", ""))
                    text_obj = c.beginText(72, 700)
                    for line in dm.get("content", "").splitlines():
                        text_obj.textLine(line)
                    c.drawText(text_obj)
                    c.showPage()
                    c.save()
                    pkg.write(pdf_file, arcname=pdf_file.name)

                if "docx" in publish_options.get("formats", []):
                    from docx import Document

                    docx_file = base_path.with_suffix(".docx")
                    doc = Document()
                    doc.add_heading(dm.get("title", ""), level=1)
                    doc.add_paragraph(dm.get("content", ""))
                    doc.save(str(docx_file))
                    pkg.write(docx_file, arcname=docx_file.name)

        return {
            "message": "Publication module published successfully",
            "pm_code": pm_code,
            "package": str(package_path),
        }
    except Exception as e:
        logger.error(f"Error publishing publication module: {str(e)}")
        raise HTTPException(500, f"Error publishing publication module: {str(e)}")


# Test endpoints for AI providers
@api_router.post("/test/text")
async def test_text_provider(text: str, task_type: str = "classify"):
    """Test text provider."""
    try:
        from .ai_providers.base import TextProcessingRequest

        text_provider = ProviderFactory.create_text_provider()
        request = TextProcessingRequest(text=text, task_type=task_type)

        if task_type == "classify":
            response = await text_provider.classify_document(request)
        elif task_type == "extract":
            response = await text_provider.extract_structured_data(request)
        elif task_type == "rewrite":
            response = await text_provider.rewrite_to_ste(request)
        else:
            raise HTTPException(400, "Invalid task type")

        return response.dict()
    except Exception as e:
        logger.error(f"Error testing text provider: {str(e)}")
        raise HTTPException(500, f"Error testing text provider: {str(e)}")


@api_router.post("/test/vision")
async def test_vision_provider(image_data: str, task_type: str = "caption"):
    """Test vision provider."""
    try:
        from .ai_providers.base import VisionProcessingRequest

        vision_provider = ProviderFactory.create_vision_provider()
        request = VisionProcessingRequest(image_data=image_data, task_type=task_type)

        if task_type == "caption":
            response = await vision_provider.generate_caption(request)
        elif task_type == "objects":
            response = await vision_provider.detect_objects(request)
        elif task_type == "hotspots":
            response = await vision_provider.generate_hotspots(request)
        else:
            raise HTTPException(400, "Invalid task type")

        return response.dict()
    except Exception as e:
        logger.error(f"Error testing vision provider: {str(e)}")
        raise HTTPException(500, f"Error testing vision provider: {str(e)}")


# Include the router in the main app
app.include_router(api_router)


# Shutdown event
@app.on_event("shutdown")
async def shutdown_db_client():
    """Close database connection on shutdown."""
    client.close()


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=8001)
